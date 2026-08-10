import io
import uuid

import pytest
from docx import Document
from fastapi.testclient import TestClient
from api import app

client = TestClient(app)


def _build_test_resume_docx(skills_text: str) -> bytes:
    document = Document()
    document.add_paragraph("Jane Doe — Software Engineer")
    document.add_paragraph(skills_text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture(scope="module")
def auth_headers():
    email = f"test-{uuid.uuid4().hex[:12]}@nextskill.dev"
    response = client.post("/auth/signup", json={"email": email, "password": "testpassword123"})
    assert response.status_code == 200, response.text
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_health():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_signup_and_login():
    email = f"test-{uuid.uuid4().hex[:12]}@nextskill.dev"
    signup = client.post("/auth/signup", json={"email": email, "password": "testpassword123"})
    assert signup.status_code == 200
    assert "access_token" in signup.json()

    login = client.post("/auth/login", json={"email": email, "password": "testpassword123"})
    assert login.status_code == 200
    assert "access_token" in login.json()


def test_signup_duplicate_email_rejected():
    email = f"test-{uuid.uuid4().hex[:12]}@nextskill.dev"
    client.post("/auth/signup", json={"email": email, "password": "testpassword123"})
    response = client.post("/auth/signup", json={"email": email, "password": "anotherpassword"})
    assert response.status_code == 409


def test_login_wrong_password_rejected():
    email = f"test-{uuid.uuid4().hex[:12]}@nextskill.dev"
    client.post("/auth/signup", json={"email": email, "password": "testpassword123"})
    response = client.post("/auth/login", json={"email": email, "password": "wrongpassword"})
    assert response.status_code == 401


def test_me_requires_auth():
    response = client.get("/auth/me")
    assert response.status_code == 401  # no Authorization header at all


def test_me_and_skills_update(auth_headers):
    me = client.get("/auth/me", headers=auth_headers)
    assert me.status_code == 200
    assert me.json()["skills"] == []

    updated = client.put("/auth/me/skills", json={"skills": ["Python", "SQL"]}, headers=auth_headers)
    assert updated.status_code == 200
    assert updated.json()["skills"] == ["Python", "SQL"]

    me_again = client.get("/auth/me", headers=auth_headers)
    assert me_again.json()["skills"] == ["Python", "SQL"]


def test_list_jobs_returns_results():
    response = client.get("/jobs?limit=5")
    assert response.status_code == 200
    data = response.json()
    assert "results" in data
    assert len(data["results"]) <= 5
    assert data["total"] > 0


def test_list_jobs_filters_by_role():
    response = client.get("/jobs?role=data scientist&limit=5")
    assert response.status_code == 200
    data = response.json()
    for job in data["results"]:
        assert "data" in job["title"].lower() or "scientist" in job["title"].lower()


def test_get_job_not_found():
    response = client.get("/jobs/999999999")
    assert response.status_code == 404


def test_top_companies():
    response = client.get("/companies/top?limit=3")
    assert response.status_code == 200
    data = response.json()
    assert len(data["results"]) <= 3
    postings = [c["postings"] for c in data["results"]]
    assert postings == sorted(postings, reverse=True)


def test_trend_for_known_skill():
    response = client.get("/trends/Python")
    assert response.status_code == 200
    data = response.json()
    assert data["skill"] == "Python"
    assert data["total_mentions"] > 0


def test_trend_for_unknown_skill():
    response = client.get("/trends/DefinitelyNotARealSkillXYZ")
    assert response.status_code == 404


def test_recommend_requires_auth():
    response = client.post(
        "/recommend",
        json={"skills": ["Python"], "target_role": "data scientist"},
    )
    assert response.status_code == 401  # no Authorization header at all


def test_recommend_returns_ranked_list(auth_headers):
    response = client.post(
        "/recommend",
        json={"skills": ["Python", "SQL"], "target_role": "data scientist"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert data["target_role"] == "data scientist"
    assert len(data["recommendations"]) > 0
    counts = [r["postings_mentioning_it"] for r in data["recommendations"]]
    assert counts == sorted(counts, reverse=True)


def test_recommend_excludes_stated_skills(auth_headers):
    response = client.post(
        "/recommend",
        json={"skills": ["Python"], "target_role": "data scientist"},
        headers=auth_headers,
    )
    data = response.json()
    recommended_names = {r["skill"].lower() for r in data["recommendations"]}
    assert "python" not in recommended_names


def test_recommend_falls_back_to_saved_skill_profile(auth_headers):
    client.put("/auth/me/skills", json={"skills": ["Python", "SQL"]}, headers=auth_headers)
    response = client.post(
        "/recommend",
        json={"target_role": "data scientist"},  # no skills in the body
        headers=auth_headers,
    )
    assert response.status_code == 200
    assert response.json()["your_skills"] == ["Python", "SQL"]


def test_recommend_evidence_includes_real_postings(auth_headers):
    response = client.post(
        "/recommend/evidence",
        json={"skills": ["Python", "SQL"], "target_role": "data scientist"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data["recommendations"]) > 0
    for rec in data["recommendations"]:
        assert "evidence" in rec
        for item in rec["evidence"]:
            assert "title" in item
            assert "company" in item


def test_related_skills_returns_sensible_results():
    response = client.get("/skills/React/related?limit=5")
    assert response.status_code == 200
    data = response.json()
    assert data["skill"] == "React"
    assert data["based_on_postings"] > 0
    assert len(data["related_skills"]) <= 5
    related_names = {r["skill"].lower() for r in data["related_skills"]}
    assert "react" not in related_names


def test_related_skills_unknown_skill():
    response = client.get("/skills/DefinitelyNotARealSkillXYZ/related")
    assert response.status_code == 404


def test_resume_upload_requires_auth():
    resume_bytes = _build_test_resume_docx("Skills: Python, SQL")
    response = client.post(
        "/auth/me/resume",
        files={"file": ("resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 401


def test_resume_upload_extracts_and_merges_skills(auth_headers):
    resume_bytes = _build_test_resume_docx("Skills: Python, TensorFlow, Docker, Kubernetes, AWS")
    response = client.post(
        "/auth/me/resume",
        files={"file": ("resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert "Python" in data["skills_found_in_resume"]
    assert "TensorFlow" in data["skills_found_in_resume"]
    assert set(data["skills_found_in_resume"]).issubset(set(data["skills"]))


def test_resume_upload_rejects_unsupported_format(auth_headers):
    response = client.post(
        "/auth/me/resume",
        files={"file": ("resume.txt", b"Skills: Python", "text/plain")},
        headers=auth_headers,
    )
    assert response.status_code == 400


def test_match_score_requires_auth():
    response = client.post("/match-score", json={"target_role": "data scientist"})
    assert response.status_code == 401


def test_match_score_higher_for_relevant_skills(auth_headers):
    relevant = client.post(
        "/match-score",
        json={"skills": ["Python", "SQL", "Pandas"], "target_role": "data scientist"},
        headers=auth_headers,
    )
    irrelevant = client.post(
        "/match-score",
        json={"skills": ["React", "CSS"], "target_role": "data scientist"},
        headers=auth_headers,
    )
    assert relevant.status_code == 200
    assert irrelevant.status_code == 200
    assert relevant.json()["sample_size"] > 0
    assert relevant.json()["match_pct"] > irrelevant.json()["match_pct"]


def test_match_score_unknown_role_returns_zero_sample(auth_headers):
    response = client.post(
        "/match-score",
        json={"skills": ["Python"], "target_role": "definitely not a tracked role xyz"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert data["sample_size"] == 0
    assert data["match_pct"] is None


def test_predict_salary_requires_auth():
    response = client.post("/predict-salary", json={"target_role": "data scientist"})
    assert response.status_code == 401


def test_predict_salary_returns_a_range(auth_headers):
    response = client.post(
        "/predict-salary",
        json={"skills": ["Python", "TensorFlow", "PyTorch", "Docker"], "target_role": "machine learning engineer"},
        headers=auth_headers,
    )
    # CI trains the model before running tests (see train_salary_model.py); locally, if no
    # model has been trained yet, this correctly returns 503 instead of crashing.
    assert response.status_code in (200, 503)
    if response.status_code == 200:
        data = response.json()
        assert data["predicted_salary_low"] <= data["predicted_salary_midpoint"] <= data["predicted_salary_high"]


def test_transition_graph_returns_nodes_and_edges():
    response = client.get("/roles/transition-graph")
    assert response.status_code == 200
    data = response.json()
    assert len(data["nodes"]) > 0
    assert all({"id", "label", "posting_count"} <= node.keys() for node in data["nodes"])
    assert all({"source", "target", "weight"} <= edge.keys() for edge in data["edges"])


def test_nearest_roles_for_known_role():
    response = client.get("/roles/data scientist/nearest?limit=3")
    assert response.status_code == 200
    data = response.json()
    assert len(data["nearest_roles"]) <= 3
    for entry in data["nearest_roles"]:
        assert "skills_you_would_need" in entry


def test_nearest_roles_for_untracked_role():
    response = client.get("/roles/underwater basket weaver/nearest")
    assert response.status_code == 404
